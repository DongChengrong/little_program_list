from docx import Document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式

#判断Unicode字符是否是数字
def is_number(uchar):
    if uchar >= u'\u0030' and uchar<=u'\u0039':
       return True
    else:
       return False

def para_change(source_document, target_document):
    for paragraph in source_document.paragraphs:

        #添加word段落
        text = paragraph.text
        p = target_document.add_paragraph(u'')
        run = p.add_run(text)

        #设置段落文字字体
        run.font.name = u'宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
        if (is_number(text)) :
            #设置字体大小
            run.font.size = Pt(8)
        else:
            run.font.size = Pt(11)
        '''
        i = i + 1
        if (i > 100):
            break
        '''

        #设置行间距
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after =Pt(0)

if __name__ == '__main__':
    #for i in range(1, 19) :
    path = 'C:\\Users\\董成荣\\Desktop\\新建文件夹\\周三多 管理学5.docx' 
        #path = path + str(i) + '.docx'
    source_document = Document(path)
    target_document = Document()
    para_change(source_document, target_document)
    target_document.save('a周三多 管理学5.docx')
